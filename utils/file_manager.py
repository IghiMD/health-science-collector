#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Modul pre vytváranie a správu súborov a priečinkovej štruktúry projektu.
Implementuje vytváranie denných a hodinových priečinkov, a generovanie DOCX a HTML súborov.
"""

import os
import logging
from datetime import datetime, timedelta
from typing import List, Dict, Optional, Tuple
# Načítanie premenných z .env súboru
from dotenv import load_dotenv
load_dotenv()  # toto načíta premenné z .env súboru do prostredia
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import ftplib

# Konfigurácia loggeru
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("logs/file_manager.log"),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger("file_manager")

# Základná konfigurácia - použitie presných názvov premenných z .env
OUTPUT_DIR = "output"
FTP_CONFIG = {
    'host': os.getenv("FTP_HOST", ""),       # srv1806-files.hstgr.io
    'user': os.getenv("FTP_USER", ""),       # u121530859.ighishomeradio.eu
    'password': os.getenv("FTP_PASS", ""),   # N...
    'directory': os.getenv("FTP_DIRECTORY", "")
}

class FileManager:
    """Trieda pre správu súborov a priečinkov projektu."""
    
    def __init__(self, output_dir: str = OUTPUT_DIR):
        """
        Inicializácia správcu súborov.
        
        Args:
            output_dir: Základný výstupný priečinok.
        """
        self.output_dir = output_dir
        self._ensure_output_dir()
    
    def _ensure_output_dir(self) -> None:
        """Zaistí existenciu základného výstupného priečinka."""
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
            logger.info(f"Vytvorený výstupný priečinok: {self.output_dir}")
    
    def create_daily_directory(self, date: Optional[datetime] = None) -> str:
        """
        Vytvorí denný priečinok s dátumom.
        
        Args:
            date: Dátum pre priečinok. Ak None, použije sa aktuálny dátum.
            
        Returns:
            Cesta k vytvorenému priečinku.
        """
        if date is None:
            date = datetime.now()
        
        date_str = date.strftime("%Y-%m-%d")
        daily_dir = os.path.join(self.output_dir, date_str)
        
        if not os.path.exists(daily_dir):
            os.makedirs(daily_dir)
            logger.info(f"Vytvorený denný priečinok: {daily_dir}")
        
        return daily_dir
    
    def create_hourly_directory(self, date: Optional[datetime] = None) -> str:
        """
        Vytvorí hodinový priečinok v rámci denného priečinka.
        
        Args:
            date: Dátum a čas pre priečinok. Ak None, použije sa aktuálny čas.
            
        Returns:
            Cesta k vytvorenému hodinovému priečinku.
        """
        if date is None:
            date = datetime.now()
        
        # Vytvorenie denného priečinka
        daily_dir = self.create_daily_directory(date)
        
        # Vytvorenie hodinového priečinka
        hour_str = date.strftime("%H-%M")
        hourly_dir = os.path.join(daily_dir, hour_str)
        
        if not os.path.exists(hourly_dir):
            os.makedirs(hourly_dir)
            logger.info(f"Vytvorený hodinový priečinok: {hourly_dir}")
        
        return hourly_dir
    
    def create_docx_file(self, articles: List[Dict], output_path: str, is_summary: bool = False) -> str:
        """
        Vytvorí DOCX súbor s článkami.
        
        Args:
            articles: Zoznam článkov na pridanie do súboru.
            output_path: Cesta k výstupnému súboru.
            is_summary: Či ide o súhrnný súbor (True) alebo hodinový súbor (False).
            
        Returns:
            Cesta k vytvorenému súboru.
        """
        # Vytvorenie nového dokumentu
        doc = docx.Document()
        
        # Nastavenie štýlu dokumentu
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Pridanie nadpisu
        title = doc.add_heading('', level=1)
        run = title.add_run('Zdravotnícke a vedecké zaujímavosti')
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 0, 128)
        
        # Pridanie podnadpisu s dátumom a časom
        current_time = datetime.now().strftime("%d.%m.%Y %H:%M")
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run(f'Generované: {current_time}')
        subtitle_run.font.size = Pt(10)
        subtitle_run.font.italic = True
        
        # Pridanie prázdneho riadku
        doc.add_paragraph()
        
        # Typ dokumentu (súhrn alebo hodinový)
        doc_type = doc.add_paragraph()
        doc_type_run = doc_type.add_run('SÚHRNNÝ DOKUMENT' if is_summary else 'HODINOVÝ PREHĽAD')
        doc_type_run.font.bold = True
        doc_type_run.font.size = Pt(12)
        
        # Pridanie prázdneho riadku
        doc.add_paragraph()
        
        # Pridanie článkov
        for i, article in enumerate(articles, 1):
            # Preložený názov článku
            article_title = doc.add_heading('', level=2)
            title_text = article.get('translated_title', article.get('title', 'Bez názvu'))
            title_run = article_title.add_run(f"{i}. {title_text}")
            title_run.font.size = Pt(14)
            
            # Metadáta o článku - zdroj
            metadata = doc.add_paragraph()
            source_url = article.get('source_url', article.get('url', 'Neznámy zdroj'))
            metadata_run = metadata.add_run(f"Zdroj: {source_url}")
            metadata_run.font.size = Pt(9)
            metadata_run.font.italic = True
            
            # Originálny text (skrátený pre prehľadnosť)
            if article.get('text'):
                orig_text_para = doc.add_paragraph()
                orig_text_para.style = 'Quote'
                orig_text = article['text']
                # Skrátenie originálneho textu
                if len(orig_text) > 500:
                    orig_text = orig_text[:500] + "..."
                orig_text_run = orig_text_para.add_run(f"Originálny text: {orig_text}")
                orig_text_run.font.size = Pt(9)
                orig_text_run.font.italic = True
            
            # Hlavný sumár - sarkastický text
            summary_para = doc.add_paragraph()
            summary_run = summary_para.add_run(article.get('summary_text', ''))
            summary_run.font.size = Pt(11)
            
            # Edukačný dovetok
            if article.get('appendix_text'):
                appendix_para = doc.add_paragraph()
                appendix_para.style = 'Quote'
                appendix_run = appendix_para.add_run(f"Dovetok: {article.get('appendix_text', '')}")
                appendix_run.font.size = Pt(10)
                appendix_run.font.italic = True
            
            # Oddeľovač medzi článkami (okrem posledného)
            if i < len(articles):
                doc.add_paragraph('------------------------------')
        
        # Uloženie dokumentu
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        logger.info(f"Vytvorený DOCX súbor: {output_path}")
        
        return output_path
    
    def create_html_file(self, articles: List[Dict], output_path: str, is_summary: bool = False) -> str:
        """
        Vytvorí HTML súbor s článkami.
        
        Args:
            articles: Zoznam článkov na pridanie do súboru.
            output_path: Cesta k výstupnému súboru.
            is_summary: Či ide o súhrnný súbor (True) alebo hodinový súbor (False).
            
        Returns:
            Cesta k vytvorenému súboru.
        """
        # HTML hlavička
        html_content = f"""<!DOCTYPE html>
<html lang="sk">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{'Súhrnný dokument' if is_summary else 'Hodinový prehľad'} - Zdravotnícke a vedecké zaujímavosti</title>
    <style>
        body {{ font-family: 'Calibri', 'Arial', sans-serif; margin: 0; padding: 20px; color: #333; }}
        .container {{ max-width: 800px; margin: 0 auto; }}
        h1 {{ color: #004080; font-size: 24px; text-align: center; }}
        .subtitle {{ text-align: center; font-style: italic; margin-bottom: 20px; font-size: 14px; }}
        .doc-type {{ font-weight: bold; text-align: center; font-size: 16px; margin: 20px 0; }}
        .article {{ margin-bottom: 30px; border-bottom: 1px solid #ddd; padding-bottom: 20px; }}
        .article:last-child {{ border-bottom: none; }}
        h2 {{ color: #333; font-size: 20px; }}
        .source {{ font-style: italic; color: #666; font-size: 12px; margin-top: 5px; }}
        .original-text {{ font-style: italic; background-color: #f9f9f9; padding: 10px; border-left: 3px solid #ddd; font-size: 13px; margin: 10px 0; }}
        .summary {{ line-height: 1.5; }}
        .appendix {{ font-style: italic; background-color: #f5f5f5; padding: 10px; border-left: 3px solid #aaa; margin-top: 10px; }}
    </style>
</head>
<body>
    <div class="container">
        <h1>Zdravotnícke a vedecké zaujímavosti</h1>
        <div class="subtitle">Generované: {datetime.now().strftime("%d.%m.%Y %H:%M")}</div>
        <div class="doc-type">{'SÚHRNNÝ DOKUMENT' if is_summary else 'HODINOVÝ PREHĽAD'}</div>
"""

        # Pridanie článkov
        for i, article in enumerate(articles, 1):
            # Preložený názov článku
            title_text = article.get('translated_title', article.get('title', 'Bez názvu'))
            source_url = article.get('source_url', article.get('url', 'Neznámy zdroj'))
            
            # Originálny text (skrátený pre prehľadnosť)
            orig_text = ""
            if article.get('text'):
                orig_text = article['text']
                # Skrátenie originálneho textu
                if len(orig_text) > 500:
                    orig_text = orig_text[:500] + "..."
            
            # HTML kód pre článok
            html_content += f"""
        <div class="article">
            <h2>{i}. {title_text}</h2>
            <div class="source">Zdroj: <a href="{source_url}" target="_blank">{source_url}</a></div>
"""
            
            if orig_text:
                html_content += f"""
            <div class="original-text">Originálny text: {orig_text}</div>
"""
            
            if article.get('summary_text'):
                html_content += f"""
            <div class="summary">{article.get('summary_text', '')}</div>
"""
            
            if article.get('appendix_text'):
                html_content += f"""
            <div class="appendix">Dovetok: {article.get('appendix_text', '')}</div>
"""
            
            html_content += """
        </div>
"""
        
        # HTML pätička
        html_content += """
    </div>
</body>
</html>
"""
        
        # Uloženie HTML súboru
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        with open(output_path, 'w', encoding='utf-8') as file:
            file.write(html_content)
        
        logger.info(f"Vytvorený HTML súbor: {output_path}")
        return output_path
    
    def create_hourly_docx(self, articles: List[Dict], date: Optional[datetime] = None) -> str:
        """
        Vytvorí hodinový DOCX súbor s článkami.
        
        Args:
            articles: Zoznam článkov na pridanie do súboru.
            date: Dátum a čas pre súbor. Ak None, použije sa aktuálny čas.
            
        Returns:
            Cesta k vytvorenému súboru.
        """
        if date is None:
            date = datetime.now()
        
        # Vytvorenie hodinového priečinka
        hourly_dir = self.create_hourly_directory(date)
        
        # Vytvorenie cesty k súboru
        date_str = date.strftime("%Y-%m-%d")
        hour_str = date.strftime("%H-%M")
        filename = f"{date_str}_{hour_str}.docx"
        file_path = os.path.join(hourly_dir, filename)
        
        # Vytvorenie súboru
        return self.create_docx_file(articles, file_path, is_summary=False)
    
    def create_hourly_html(self, articles: List[Dict], date: Optional[datetime] = None) -> str:
        """
        Vytvorí hodinový HTML súbor s článkami.
        
        Args:
            articles: Zoznam článkov na pridanie do súboru.
            date: Dátum a čas pre súbor. Ak None, použije sa aktuálny čas.
            
        Returns:
            Cesta k vytvorenému súboru.
        """
        if date is None:
            date = datetime.now()
        
        # Vytvorenie hodinového priečinka
        hourly_dir = self.create_hourly_directory(date)
        
        # Vytvorenie cesty k súboru
        date_str = date.strftime("%Y-%m-%d")
        hour_str = date.strftime("%H-%M")
        filename = f"{date_str}_{hour_str}.html"
        file_path = os.path.join(hourly_dir, filename)
        
        # Vytvorenie súboru
        return self.create_html_file(articles, file_path, is_summary=False)
    
    def create_or_update_summary_docx(self, articles: List[Dict], date: Optional[datetime] = None) -> str:
        """
        Vytvorí alebo aktualizuje súhrnný DOCX súbor pre daný deň.
        
        Args:
            articles: Zoznam článkov na pridanie do súboru.
            date: Dátum pre súbor. Ak None, použije sa aktuálny dátum.
            
        Returns:
            Cesta k vytvorenému/aktualizovanému súboru.
        """
        if date is None:
            date = datetime.now()
        
        # Vytvorenie denného priečinka
        daily_dir = self.create_daily_directory(date)
        
        # Vytvorenie cesty k súboru
        date_str = date.strftime("%Y-%m-%d")
        filename = f"{date_str}_summary.docx"
        file_path = os.path.join(daily_dir, filename)
        
        # Vytvorenie súboru
        return self.create_docx_file(articles, file_path, is_summary=True)
    
    def create_or_update_summary_html(self, articles: List[Dict], date: Optional[datetime] = None) -> str:
        """
        Vytvorí alebo aktualizuje súhrnný HTML súbor pre daný deň.
        
        Args:
            articles: Zoznam článkov na pridanie do súboru.
            date: Dátum pre súbor. Ak None, použije sa aktuálny dátum.
            
        Returns:
            Cesta k vytvorenému/aktualizovanému súboru.
        """
        if date is None:
            date = datetime.now()
        
        # Vytvorenie denného priečinka
        daily_dir = self.create_daily_directory(date)
        
        # Vytvorenie cesty k súboru
        date_str = date.strftime("%Y-%m-%d")
        filename = f"{date_str}_summary.html"
        file_path = os.path.join(daily_dir, filename)
        
        # Vytvorenie súboru
        return self.create_html_file(articles, file_path, is_summary=True)
    
    def create_hourly_files(self, articles: List[Dict], date: Optional[datetime] = None) -> Tuple[str, str]:
        """
        Vytvorí hodinové DOCX a HTML súbory s článkami.
        
        Args:
            articles: Zoznam článkov na pridanie do súborov.
            date: Dátum a čas pre súbory. Ak None, použije sa aktuálny čas.
            
        Returns:
            Tuple (docx_path, html_path) s cestami k vytvoreným súborom.
        """
        docx_path = self.create_hourly_docx(articles, date)
        html_path = self.create_hourly_html(articles, date)
        return docx_path, html_path
    
    def create_or_update_summary_files(self, articles: List[Dict], date: Optional[datetime] = None) -> Tuple[str, str]:
        """
        Vytvorí alebo aktualizuje súhrnné DOCX a HTML súbory pre daný deň.
        
        Args:
            articles: Zoznam článkov na pridanie do súborov.
            date: Dátum pre súbory. Ak None, použije sa aktuálny dátum.
            
        Returns:
            Tuple (docx_path, html_path) s cestami k vytvoreným/aktualizovaným súborom.
        """
        docx_path = self.create_or_update_summary_docx(articles, date)
        html_path = self.create_or_update_summary_html(articles, date)
        return docx_path, html_path
    
    def upload_to_ftp(self, local_path: str) -> bool:
        """
        Nahrá súbor na FTP server.
        
        Args:
            local_path: Lokálna cesta k súboru.
            
        Returns:
            True ak upload prebehol úspešne, inak False.
        """
        if not all([FTP_CONFIG['host'], FTP_CONFIG['user'], FTP_CONFIG['password']]):
            logger.warning("FTP konfigurácia nie je kompletná. Preskakujem upload.")
            return False
        
        try:
            # Pripojenie na FTP server
            ftp = ftplib.FTP(FTP_CONFIG['host'])
            ftp.login(FTP_CONFIG['user'], FTP_CONFIG['password'])
            
            # Nastavenie adresára
            if FTP_CONFIG['directory']:
                try:
                    ftp.cwd(FTP_CONFIG['directory'])
                except ftplib.error_perm:
                    # Adresár neexistuje, skúsime ho vytvoriť
                    ftp.mkd(FTP_CONFIG['directory'])
                    ftp.cwd(FTP_CONFIG['directory'])
            
            # Vytvorenie potrebnej adresárovej štruktúry na FTP
            relative_path = os.path.relpath(local_path, self.output_dir)
            directories = os.path.dirname(relative_path).split(os.sep)
            
            for directory in directories:
                if directory:
                    try:
                        ftp.cwd(directory)
                    except ftplib.error_perm:
                        ftp.mkd(directory)
                        ftp.cwd(directory)
            
            # Upload súboru
            with open(local_path, 'rb') as file:
                ftp.storbinary(f'STOR {os.path.basename(local_path)}', file)
            
            # Ukončenie FTP spojenia
            ftp.quit()
            
            logger.info(f"Súbor {local_path} úspešne nahraný na FTP server.")
            return True
        
        except Exception as e:
            logger.error(f"Chyba pri FTP uploade súboru {local_path}: {str(e)}")
            return False
    
    def upload_multiple_to_ftp(self, local_paths: List[str]) -> List[bool]:
        """
        Nahrá viacero súborov na FTP server.
        
        Args:
            local_paths: Zoznam lokálnych ciest k súborom.
            
        Returns:
            Zoznam výsledkov uploadov (True pre úspešný upload, False pre neúspešný).
        """
        results = []
        for path in local_paths:
            result = self.upload_to_ftp(path)
            results.append(result)
        return results

# Príklad použitia
if __name__ == "__main__":
    manager = FileManager()
    
    # Test vytvorenia priečinkov
    daily_dir = manager.create_daily_directory()
    hourly_dir = manager.create_hourly_directory()
    
    # Testovací článok
    test_article = {
        'title': 'Testovací článok',
        'url': 'https://example.com/article1',
        'summary_text': 'Toto je testovací sumár článku. Obsahuje niekoľko viet, ktoré simulujú obsah skutočného článku o zdraví a vede.',
        'appendix_text': 'Vedeli ste, že pravidelná konzumácia zeleniny znižuje riziko srdcových ochorení o 25%?'
    }
    
    # Test vytvorenia DOCX a HTML súborov
    docx_file, html_file = manager.create_hourly_files([test_article])
    summary_docx, summary_html = manager.create_or_update_summary_files([test_article])
    
    print(f"Vytvorené súbory:\n- {docx_file}\n- {html_file}\n- {summary_docx}\n- {summary_html}")